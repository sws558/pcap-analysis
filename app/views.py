# from posixpath import abspath
import multiprocessing
import hashlib
import os
import time
from datetime import datetime
import zipfile
import bcrypt
# from multiprocessing import Process, Lock

from flask import (Response,flash, redirect, render_template, request,jsonify,
                   send_from_directory, url_for,session)
from scapy.all import *

from scapy.all import rdpcap

from app import app
from app.models import predict_csv
from app.utils.data_extract import (client_info, mail_data, sen_data, telnet_ftp_data, web_data)
from app.utils.except_info import exception_warning
from app.utils.file_extract import all_files, ftp_file, mail_file, web_file
from app.utils.flow_analyzer import (data_flow, data_in_out_ip, get_host_ip, most_flow_statistic, proto_flow, time_flow)
from app.utils.ipmap_tools import get_geo, get_ipmap, getmyip
from app.utils.pcap_decode import PcapDecode
from app.utils.pcap_filter import get_all_pcap, proto_filter, showdata_from_id, get_time_pcap
#调用 line185
from app.utils.proto_analyzer import (common_proto_statistic, dns_statistic, http_statistic, most_proto_statistic, pcap_len_statistic)
from .models import *

from flask_jwt_extended import create_access_token, create_refresh_token, jwt_required, get_jwt_identity
import shutil

from flask_paginate import Pagination

from .forms import Upload
from .utils.upload_tools import (allowed_file, get_filetype, pcapng_to_pcap, random_name)
#额外
from docx.shared import Inches
from flask import Flask, render_template, send_file
from scapy.all import rdpcap
import openpyxl
from io import BytesIO
from flask import make_response
from flask import g
import xlwt
import pyx
from collections import Counter
from docx import Document
from matplotlib.figure import Figure
from matplotlib.backends.backend_agg import FigureCanvasAgg as FigureCanvas
from .models import *
from flask_jwt_extended import create_access_token, create_refresh_token, jwt_required, get_jwt_identity
from XuNet.xu_net import XuNet
from XuNet.stego_detector import StegoDetector
import torch
from PIL import Image
import numpy as np
from torchvision import transforms
import torch.nn as nn   
import pandas as pd  
# 导入函数到模板中
app.jinja_env.globals["enumerate"] = enumerate
# 导入函数到模板中
app.jinja_env.globals["enumerate"] = enumerate
PDF_NAME = None
# 全局变量
PCAP_NAME = ""  # 上传文件名
PD = PcapDecode()  # 解析器
PCAPS = None  # 数据包
SORT = []#压缩包文件按时间顺序排列
flow_df = None
flow_df3 = None
TIMELY = False
pcaps = None
# 登录
@app.route("/", methods=["POST", "GET"])

def login():
    return render_template('./login/login.html')


@app.route("/ac_login", methods=["POST", "GET"])

def ac_login():
    if request.method == "GET":
        return render_template("./login/login.html")

    elif request.method == "POST":
        data = request.get_json()
        account = data.get('account')
        password = data.get('password')


        # account = request.form.get("account", type=str, default=None)
        # password = request.form.get("password", type=str, default=None)

        finduser = User.query.filter(User.account==account)

        if finduser.count()==0:
            # return render_template("./login/login.html", msg="用户名或密码错误")
            return jsonify(code="500", msg="用户名或密码错误")
        else:
            if(bcrypt.checkpw(password.encode('utf-8'), finduser[0].hashed_password)):

                # todo 登录成功信息保存
                access_token = create_access_token(identity=account)
                refresh_token = create_refresh_token(identity=account)

                # session['access_token'] = access_token 

                # session.permanent = False  # 设置为非永久 session，过期后需要重新登录  
                # session.modified = True  # 确保 session 数据被修改并重新存储  
                # expires_in = 600  # 过期时间为 10 分钟，单位为秒  
                # session['expires_at'] = time.time() + expires_in  # 存储过期时间戳  

                return jsonify(code="200", access_token=access_token, refresh_token=refresh_token)

            else:
                # return render_template("./login/login.html", msg="用户名或密码错误")
                return jsonify(code="500", msg="用户名或密码错误")


# 注册

@app.route("/ac_register", methods=["POST", "GET"])

def ac_register():
    if request.method == "GET":
        return render_template("./login/login.html")

    elif request.method == "POST":

        # data = request.get_json()
        # account = data.get('account')
        # password = data.get('password')

        account = request.form.get("r_account", type=str, default=None)
        password = request.form.get("r_password", type=str, default=None)
        hashed_password = bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt())

        finduser = User.query.filter(User.account==account)


        if finduser.count()==0:
            u = User()
            u.account = account
            u.hashed_password = hashed_password
            db.session.add(u)
            db.session.commit()

            # todo 直接登录token存储

            # todo 登录成功信息保存
            # access_token = create_access_token(identity=account)
            # refresh_token = create_refresh_token(identity=account)


            # session['access_token'] = access_token 

            # session.permanent = False  # 设置为非永久 session，过期后需要重新登录  
            # session.modified = True  # 确保 session 数据被修改并重新存储  
            # expires_in = 600  # 过期时间为 10 分钟，单位为秒  
            # session['expires_at'] = time.time() + expires_in  # 存储过期时间戳  


            # return jsonify(access_token=access_token, refresh_token=refresh_token)

            # return render_template('./home/index.html')
            return render_template("./login/login.html", msg="注册成功")
        
        else:
            return render_template("./login/login.html", msg="用户名已被注册")


# --------------------------------------------------------首页，上传------------
# 首页

@app.route("/index/", methods=["POST", "GET"])
@jwt_required()

def index():
    current_user = get_jwt_identity()
    return render_template("./home/index.html"), 200, {'Content-Type': 'text/html'}
def upload():
    filepath = app.config["UPLOAD_FOLDER"]
    filepath = os.path.abspath(filepath) + "/"
    upload = Upload()
    if request.method == "GET":
        return render_template("./upload/upload.html")

    elif request.method == "POST":

        pcap = upload.pcap.data
        if upload.validate_on_submit():
            pcapname = pcap.filename
            if allowed_file(pcapname):
                name1 = random_name()
                name2 = get_filetype(pcapname)

                global PCAP_NAME, PCAPS, PCAP_PATH, PCAPS_NUMBER, PCAPS_DIC
                PCAP_NAME = name1 + name2
                try:
                    PCAP_PATH = os.path.join(filepath, PCAP_NAME)
                    pcap.save(PCAP_PATH)
                    if PCAP_NAME.endswith(".pcapng"):
                        pcapng_to_pcap(PCAP_PATH)
                        PCAP_NAME = PCAP_NAME.replace(".pcapng", ".pcap")
                        PCAP_PATH = os.path.join(filepath, PCAP_NAME)

                    PCAPS = rdpcap(PCAP_PATH)
                    
                    flash("恭喜你,上传成功！")
                    return render_template("./upload/upload.html")
                except Exception as e:
                    flash("上传错误,错误信息:" + str(e))
                    return render_template("./upload/upload.html")
            else:
                flash("上传失败,请上传允许的数据包格式!")
                return render_template("./upload/upload.html")
        else:
            return render_template("./upload/upload.html")



# -----------------------实时分析---------------------

@app.route("/timelyanalysis", methods=["POST", "GET"])

 
        
def timelyanalysis():
    
    if request.method == "GET":
        
        return render_template("./upload/upload.html", title="实时分析")
    
    elif request.method == "POST":
        
        global PCAPS, PCAP_PATH, PCAP_NAME, TIMELY
        now_date = "{}".format(time.strftime("%Y%m%d", time.localtime(time.time()-60)))
        PCAP_NAME = "{}.pcap".format(time.strftime("%Y%m%d-%H:%M", time.localtime(time.time()-60)))
        PCAP_PATH = os.path.join(app.config["UPLOAD_FOLDER"], now_date, PCAP_NAME)
        # while not os.path.exists(PCAP_PATH):
            # continue
        PCAPS = rdpcap(PCAP_PATH)
        if len(PCAPS) > 1000:
            PCAPS = PCAPS[-1000:]
        flash("读取成功")
        TIMELY = True
        flash("Begining real-time sniff packages")
        
        return render_template("./upload/upload.html", title="实时分析")


# 数据包上传
@app.route("/upload/", methods=["POST", "GET"])
def upload():
    global TIMELY
    TIMELY = False
    filepath = app.config["UPLOAD_FOLDER"]
    filepath = os.path.abspath(filepath) + "/"
    zippath = app.config["ZIP_PATH"]
    num = app.config["NUM"]
    zippath = os.path.abspath(zippath)
    upload = Upload()
    if request.method == "GET":
        return render_template("./upload/upload.html")

    elif request.method == "POST":
        pcap = upload.pcap.data
        if upload.validate_on_submit():
            pcapname = pcap.filename
            if allowed_file(pcapname):
                name1 = random_name()
                name2 = get_filetype(pcapname)
                file_info_list=[]
                global PCAP_NAME, PCAPS, PCAP_PATH, SORT
                PCAP_NAME = name1 + name2

                try:
                    # PCAP_PATH = os.path.join(filepath, PCAP_NAME)
                    # pcap.save(PCAP_PATH)
                    if PCAP_NAME.endswith(".pcapng"):
                        pcapng_to_pcap(PCAP_PATH)
                        PCAP_NAME = PCAP_NAME.replace(".pcapng", ".pcap")
                        PCAP_PATH = os.path.join(filepath, PCAP_NAME)
                    PCAP_PATH = os.path.join(filepath, PCAP_NAME)
                    pcap.save(PCAP_PATH)
                    PCAPS = rdpcap(PCAP_PATH)
                    with zipfile.ZipFile(zippath, 'a') as zipf:
                        zipf.write(PCAP_PATH, arcname=PCAP_NAME)
                    with zipfile.ZipFile(zippath, 'r') as zipf:
                        for file_info in zipf.infolist():
                            file_name = file_info.filename
                            create_time = file_info.date_time
                            add_time = datetime(*create_time)
                            file_info_list.append((file_name, add_time))
                    SORT = sorted(file_info_list, key=lambda x: x[1])
                    le = len(SORT)
                    if le >= 2:
                        delete = SORT[le-2][0]
                        path = os.path.join(filepath, delete)
                        os.remove(path)
                    if le >= num:
                        file_to_zip=[]
                        file = []
                        #解压缩
                        with zipfile.ZipFile(zippath, 'r') as zip:
                            zip.extractall(filepath)
                        os.remove(zippath)
                        #获取文件列表
                        file_to_zip = SORT[-num:le]
                        #获取文件名
                        for i in range(num):
                            file = file_to_zip[i][0]
                            # print(i,file,type(file))
                            path = os.path.join(filepath, file)
                            # print(path)
                            with zipfile.ZipFile(zippath, 'a') as zipf:
                                zipf.write(path, arcname=file)
                        # with zipfile.ZipFile(zippath, 'r') as zip_ref:
                        #     # 获取zip文件中的文件名列表
                        #     file_names = zip_ref.namelist()
                        #     # 打印文件名列表
                        #     print("压缩文件中的文件名:")
                        #     for file_name in file_names:
                        #         print(file_name)
                    file_list = os.listdir(filepath)
                    for file_name in file_list:
                        if file_name.endswith('.pcap') and file_name!= PCAP_NAME:
                            file_path = os.path.join(filepath, file_name)
                            os.remove(file_path)
                    flash("恭喜你,上传成功！")
                    return render_template("./upload/upload.html")
                except Exception as e:
                    flash("上传错误,错误信息:" + str(e))
                    return render_template("./upload/upload.html")
            else:
                flash("上传失败,请上传允许的数据包格式!")
                return render_template("./upload/upload.html")
        else:
            return render_template("./upload/upload.html")
#额外
# -------------------------------------------下载数据--------------------------
@app.route("/download/", methods=["POST", "GET"])
def download():
    # 读取 pcap 文件
    global pcaps
    packets = pcaps

    # 创建一个存储数据包信息的列表
    packet_list = []
    packet = []
    # 提取数据包信息并添加到列表
    for i in range(len(packets)):
    # for packet in packets:
        packet = packets[i+1]
        time = packet["time"]
        source = packet["Source"]
        destination = packet["Destination"]
        protocol = packet["Procotol"]
        length = packet["len"]
        info = packet["info"]

        packet_list.append({'Time': time,
                            'Source': source,
                            'Destination': destination,
                            'Protocol': protocol,
                            'Length': length,
                            'Info': info})

    # 创建一个 Workbook 和一个工作表
    wb = openpyxl.Workbook()
    ws = wb.active

    # 将数据写入工作表
    headers = ['Time', 'Source', 'Destination', 'Protocol', 'Length', 'Info']
    ws.append(headers)

    for packet_info in packet_list:
        ws.append([packet_info[header] for header in headers])

    # 将 Workbook 写入 BytesIO 缓冲区
    excel_buffer = BytesIO()
    wb.save(excel_buffer)
    excel_buffer.seek(0)             

    # 将 Excel 文件作为响应返回
    return send_file(excel_buffer, download_name='pcap_data.xlsx', as_attachment=True)

#额外
#-----------------------------协议分析界面生成分析报告----------------
#常见协议统计
def generate_proto_bar_chart(data, chart_title, color):
    fig, ax = plt.subplots()
    x_labels = ["IP", "IPv6", "TCP", "UDP", "ARP", "ICMP", "DNS", "HTTP", "HTTPS", "Others"]
    values = data
    ax.bar(x_labels, values, color=color)
    ax.set_xlabel('协议类型')
    ax.set_ylabel('数据包个数')
    ax.set_title(chart_title)
    return fig

def add_proto_bar_chart_to_word(doc, chart_data, chart_title, color):
    # 生成协议统计条形图
    proto_chart = generate_proto_bar_chart(chart_data, chart_title, color)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    proto_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading(chart_title, level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))
    
#数据包长度饼状图
def generate_pcap_len_pie_chart_1(pcap_len_data, chart_title):
    fig, ax = plt.subplots()
    labels = ['0-300', '301-600', '601-900', '901-1200', '1201-1500']
    sizes = [pcap_len_data[label] for label in labels]
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')  # 保证饼图是一个圆
    ax.set_title(chart_title)
    return fig

def add_pcap_len_pie_chart_to_word_1(doc, pcap_len_data, chart_title):
    # 生成数据包长度饼状图
    pcap_len_chart = generate_pcap_len_pie_chart_1(pcap_len_data, chart_title)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    pcap_len_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('数据包长度统计', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))

@app.route("/downloadp/", methods=["POST", "GET"])
def downloadp():
    global PCAPS, PD
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        data_dict = common_proto_statistic(PCAPS)
        pcap_len_dict = pcap_len_statistic(PCAPS)
        pcap_count_dict = most_proto_statistic(PCAPS, PD)
        http_dict = http_statistic(PCAPS)
        http_dict = sorted(http_dict.items(),
                           key=lambda d: d[1], reverse=False)
        http_key_list = list()
        http_value_list = list()
        for key, value in http_dict:
            http_key_list.append(key)
            http_value_list.append(value)
        dns_dict = dns_statistic(PCAPS)
        dns_dict = sorted(dns_dict.items(), key=lambda d: d[1], reverse=False)
        dns_key_list = list()
        dns_value_list = list()
        for key, value in dns_dict:
            dns_key_list.append(key.decode("utf-8"))
            dns_value_list.append(value)
                # 将文档保存到 BytesIO 对象中
        
        #------------------------
        pcaps = get_all_pcap(PCAPS, PD)
        total_entries = len(pcaps)
        
        #统计每个协议出现的次数
        protocol_counts = Counter(entry['Procotol'] for entry in pcaps.values())
        print("协议出现次数:", protocol_counts)
    
        
        # 找出最多的Procotol及其计数
        protocol_counts = Counter(entry.get('Procotol', 'Unknown') for entry in pcaps.values())
        most_common_protocol, count = protocol_counts.most_common(1)[0]
        
        
        # 统计 HTTP 和 HTTPS 访问次数
        http_counter = Counter(entry.get('Procotol', 'Unknown') for entry in pcaps.values() if entry.get('Procotol', 'Unknown') == 'HTTP')
        https_counter = Counter(entry.get('Procotol', 'Unknown') for entry in pcaps.values() if entry.get('Procotol', 'Unknown') == 'HTTPS')
        
        #------------------------
        doc = Document()
        doc.add_heading('协议分析统计', level=1)


        
        add_proto_bar_chart_to_word(doc, data_dict.values(), 'Common Protocol Statistics', '#87cefa')
        
        add_pcap_len_pie_chart_to_word_1(doc, pcap_len_dict, 'Packet length statistics')


        #1111111111111111111111111111
        #添加
        doc.add_heading("条目数", level=2)
       # doc.add_picture(plot_path, width=Inches(4), height=Inches(3))


        # 将统计信息添加到文档
        doc.add_paragraph(f"总条目数: {total_entries}")

        doc.add_heading('协议出现次数', level=2)
        for protocol, count in protocol_counts.items():
            doc.add_paragraph(f"{protocol}: {count}")

        doc.add_heading('最多的协议', level=2)
        doc.add_paragraph(f"最多的Procotol: {most_common_protocol}")
        doc.add_paragraph(f"出现次数: {count}")

        doc.add_heading('HTTP 和 HTTPS 访问次数', level=2)
        doc.add_paragraph(f"HTTP 访问次数: {http_counter['HTTP']}")
        doc.add_paragraph(f"HTTPS 访问次数: {https_counter['HTTPS']}")
        #1111111111111111111111111111
        
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        
        # 将 BytesIO 对象的内容发送给用户
        doc_buffer.seek(0)
        return send_file(doc_buffer, as_attachment=True, download_name='statistics_report.docx')
    
    

FILTER_TYPE_DB = None
VALUE_DB = None
STARTTIME_DB = None
ENDTIME_DB = None
# -------------------------------------------数据分析--------------------------
@app.route("/database/", methods=["POST", "GET"])
def basedata():
    """
    基础数据解析
    """
    global PCAPS, PD
    global pcaps
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        if TIMELY:
            global PCAP_NAME, PCAP_PATH
            now_date = "{}".format(time.strftime("%Y%m%d", time.localtime(time.time()-60)))
            PCAP_NAME = "{}.pcap".format(time.strftime("%Y%m%d-%H:%M", time.localtime(time.time()-60)))
            PCAP_PATH = os.path.join(app.config["UPLOAD_FOLDER"], now_date, PCAP_NAME)
            # while not os.path.exists(PCAP_PATH):
                # continue
            PCAPS = rdpcap(PCAP_PATH)
            if len(PCAPS) > 1000:
                PCAPS = PCAPS[-1000:]
        # 将筛选的type和value通过表单获取
        filter_type = request.form.get("filter_type", type=str, default=None)
        value = request.form.get("value", type=str, default=None)
        starttime = request.form.get("starttime", type=str, default=None)        
        endtime = request.form.get("endtime", type=str, default=None)
        page = request.args.get("page", type=int, default=None)

        global FILTER_TYPE_DB, VALUE_DB, STARTTIME_DB, ENDTIME_DB

        if filter_type == None:
            value = VALUE_DB
            starttime = STARTTIME_DB
            endtime = ENDTIME_DB
            filter_type = FILTER_TYPE_DB
        else:
            VALUE_DB = value
            STARTTIME_DB = starttime
            ENDTIME_DB = endtime
            FILTER_TYPE_DB = filter_type

        # starttime = datetime.strptime(starttime, "%Y-%m-%dT%H:%M")
        # endtime = datetime.strptime(endtime, "%Y-%m-%dT%H:%M")

        # 如果有选择，通过选择来获取值
        if filter_type and value:
            pcaps = proto_filter(filter_type, value, PCAPS, PD)
            # start_time = request.form.get('starttime', type=str, default=None)
            # end_time = request.form.get('endtime')     
            # if 1 and 1:#默认添加时分秒
            #     pcaps = get_time_pcap(pcaps, "2016-12-13 19:52:44", "2016-12-13 19:52:51", PD)
        #默认显示所有的协议数据
        else:
            pcaps = get_all_pcap(PCAPS, PD)
        if starttime and endtime:#默认添加时分秒
            starttime = datetime.strptime(starttime, "%Y-%m-%dT%H:%M")
            endtime = datetime.strptime(endtime, "%Y-%m-%dT%H:%M")
            pcaps = get_time_pcap(pcaps, starttime, endtime, PD)
        
        if page == None:
            page = 1
        per_page = 15


        # 转换为支持分页的数据结构
        pagination = Pagination(page=page, total=len(pcaps), per_page=per_page)
        pcaps_list = list(pcaps.values())  # 将字典的值转换为列表
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        return render_template("./dataanalyzer/basedata.html", pcaps=pcaps_paginated, pagination=pagination)


FILTER_TYPE_LLM = None
VALUE_LLM = None

# 大模型
@app.route("/LLM/", methods=["POST", "GET"])
def LLM():
    """
    基础数据解析
    """
    global PCAPS, PD
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        if TIMELY:
            global PCAP_NAME, PCAP_PATH
            now_date = "{}".format(time.strftime("%Y%m%d", time.localtime(time.time()-60)))
            PCAP_NAME = "{}.pcap".format(time.strftime("%Y%m%d-%H:%M", time.localtime(time.time()-60)))
            PCAP_PATH = os.path.join(app.config["UPLOAD_FOLDER"], now_date, PCAP_NAME)
            # while not os.path.exists(PCAP_PATH):
                # continue
            PCAPS = rdpcap(PCAP_PATH)
            if len(PCAPS) > 1000:
                PCAPS = PCAPS[-1000:]
        # 将筛选的type和value通过表单获取
        filter_type = request.form.get("filter_type", type=str, default=None)
        value = request.form.get("value", type=str, default=None)
        page = request.args.get("page", type=int, default=None)

        global FILTER_TYPE_LLM, VALUE_LLM

        if filter_type == None:
            value = VALUE_LLM
            filter_type = FILTER_TYPE_LLM
        else:
            VALUE_LLM = value
            FILTER_TYPE_LLM = filter_type
        
        # 如果有选择，通过选择来获取值
        if filter_type and value:
            pcaps = proto_filter(filter_type, value, PCAPS, PD)
        # 默认显示所有的协议数据
        else:
            pcaps = get_all_pcap(PCAPS, PD)
        
        if page == None:
            page = 1
        per_page = 15


        # 转换为支持分页的数据结构
        pagination = Pagination(page=page, total=len(pcaps), per_page=per_page)
        pcaps_list = list(pcaps.values())  # 将字典的值转换为列表
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        return render_template("./dataanalyzer/LLM.html", pcaps=pcaps_paginated, pagination=pagination)

PDF_NAME = ""
# 详细数据

@app.route("/datashow/", methods=["POST", "GET"])
def datashow():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        global PDF_NAME
        dataid = request.args.get("id")
        dataid = int(dataid) - 1
        data = showdata_from_id(PCAPS, dataid)
        PDF_NAME = random_name() + ".pdf"
        # try:
        # PCAPS[dataid].pdfdump(app.config["PDF_FOLDER"] + PDF_NAME)
        # except:
        #     pass
        return data


# 将数据包保存为pdf


@app.route("/savepdf/", methods=["POST", "GET"])
def savepdf():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        global PDF_NAME
        # pdf = app.config["PDF_FOLDER"]
        # path = os.path.join(pdf, PDF_NAME)
        # return send_file(path, as_attachment=True)
        return send_from_directory(
            app.config["PDF_FOLDER"], PDF_NAME, as_attachment=True
        )


# 协议分析
@app.route("/protoanalyzer/", methods=["POST", "GET"])
def protoanalyzer():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        data_dict = common_proto_statistic(PCAPS)
        pcap_len_dict = pcap_len_statistic(PCAPS)
        pcap_count_dict = most_proto_statistic(PCAPS, PD)
        http_dict = http_statistic(PCAPS)
        http_dict = sorted(http_dict.items(),
                           key=lambda d: d[1], reverse=False)
        http_key_list = list()
        http_value_list = list()
        for key, value in http_dict:
            http_key_list.append(key)
            http_value_list.append(value)
        dns_dict = dns_statistic(PCAPS)
        dns_dict = sorted(dns_dict.items(), key=lambda d: d[1], reverse=False)
        dns_key_list = list()
        dns_value_list = list()
        for key, value in dns_dict:
            dns_key_list.append(key.decode("utf-8"))
            dns_value_list.append(value)
        return render_template(
            "./dataanalyzer/protoanalyzer.html",
            data=list(data_dict.values()),
            pcap_len=pcap_len_dict,
            pcap_keys=list(pcap_count_dict.keys()),
            http_key=http_key_list,
            http_value=http_value_list,
            dns_key=dns_key_list,
            dns_value=dns_value_list,
            pcap_count=pcap_count_dict,
        )


# 流量分析


@app.route("/flowanalyzer/", methods=["POST", "GET"])
def flowanalyzer():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        time_flow_dict = time_flow(PCAPS)
        host_ip = get_host_ip(PCAPS)
        data_flow_dict = data_flow(PCAPS, host_ip)
        data_ip_dict = data_in_out_ip(PCAPS, host_ip)
        proto_flow_dict = proto_flow(PCAPS)
        most_flow_dict = most_flow_statistic(PCAPS, PD)
        most_flow_dict = sorted(
            most_flow_dict.items(), key=lambda d: d[1], reverse=True
        )
        if len(most_flow_dict) > 10:
            most_flow_dict = most_flow_dict[0:10]
        most_flow_key = list()
        for key, value in most_flow_dict:
            most_flow_key.append(key)
        return render_template(
            "./dataanalyzer/flowanalyzer.html",
            time_flow_keys=list(time_flow_dict.keys()),
            time_flow_values=list(time_flow_dict.values()),
            data_flow=data_flow_dict,
            ip_flow=data_ip_dict,
            proto_flow=list(proto_flow_dict.values()),
            most_flow_key=most_flow_key,
            most_flow_dict=most_flow_dict,
        )

#downloadf使用过的功能函数
#生成时间流量图函数
def generate_time_flow_chart(time_flow_dict):
    # 生成时间流量曲线图
    fig, ax = plt.subplots()
    time_points = list(time_flow_dict.keys())
    values = list(time_flow_dict.values())
    ax.plot(time_points, values, label='time discharge curve')
    ax.set_xlabel('Relative time')
    ax.set_ylabel('Packet Bytes')
    ax.legend()
    return fig
#生成饼状图函数
def generate_pie_chart(data):
    fig, ax = plt.subplots()
    labels = list(data.keys())
    sizes = list(data.values())
    ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')  # 保证饼图是一个圆
    return fig

def add_pie_chart_to_word(doc, chart_data, chart_title):
    # 生成饼状图
    pie_chart = generate_pie_chart(chart_data)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    pie_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading(chart_title, level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))
#流入生成柱状图
def generate_bar_chart(data, chart_title):
    fig, ax = plt.subplots()
    x_labels = data['in_keyp']
    values = data['in_packet']
    ax.barh(x_labels, values, color='#6495ed')  # 使用水平柱状图
    ax.set_xlabel('Number of data packets')
    ax.set_ylabel('Inflow IP')
    ax.set_title(chart_title)
    return fig

def add_bar_chart_to_word(doc, chart_data, chart_title):
    # 生成柱状图
    bar_chart = generate_bar_chart(chart_data, chart_title)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    bar_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('流入IP流量数据包个数图', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))

#流出生成柱状图
def generate_bar_chart_1(data, chart_title):
    fig, ax = plt.subplots()
    x_labels = data['out_keyp']
    values = data['out_packet']
    ax.barh(x_labels, values, color='#6495ed')  # 使用水平柱状图
    ax.set_xlabel('Number of data packets')
    ax.set_ylabel('out IP')
    ax.set_title(chart_title)
    return fig

def add_bar_chart_to_word_1(doc, chart_data, chart_title):
    # 生成柱状图
    bar_chart = generate_bar_chart_1(chart_data, chart_title)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    bar_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('流出IP流量数据包个数图', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))
#流入IP总流量图
def generate_bar_chart_2(data, chart_title):
    fig, ax = plt.subplots()
    x_labels = data['in_keyl']
    values = data['in_len']
    ax.barh(x_labels, values, color='#6495ed')  # 使用水平柱状图
    ax.set_xlabel('Total packet traffic')
    ax.set_ylabel('Inflow IP')
    ax.set_title(chart_title)
    return fig

def add_bar_chart_to_word_2(doc, chart_data, chart_title):
    # 生成柱状图
    bar_chart = generate_bar_chart_2(chart_data, chart_title)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    bar_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('流入IP总流量图', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))
#流出IP总流量图
def generate_bar_chart_3(data, chart_title):
    fig, ax = plt.subplots()
    x_labels = data['out_keyl']
    values = data['out_len']
    ax.barh(x_labels, values, color='#6495ed')  # 使用水平柱状图
    ax.set_xlabel('Total packet traffic')
    ax.set_ylabel('Inflow IP')
    ax.set_title(chart_title)
    return fig

def add_bar_chart_to_word_3(doc, chart_data, chart_title):
    # 生成柱状图
    bar_chart = generate_bar_chart(chart_data, chart_title)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    bar_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('流出IP总流量图', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))
#常见协议统计
def generate_proto_bar_chart(data, chart_title, color):
    fig, ax = plt.subplots()
    x_labels = ["IP", "IPv6", "TCP", "UDP", "ARP", "ICMP", "DNS", "HTTP", "HTTPS", "Others"]
    values = data
    ax.bar(x_labels, values, color=color)
    ax.set_xlabel('Protocol')
    ax.set_ylabel('Number of data packets')
    ax.set_title('Common Protocol Statistics')
    return fig

def add_proto_bar_chart_to_word(doc, chart_data, chart_title, color):
    # 生成协议统计条形图
    proto_chart = generate_proto_bar_chart(chart_data, chart_title, color)

    # 将图表保存到 BytesIO 对象中
    chart_buffer = BytesIO()
    proto_chart.savefig(chart_buffer, format='png')
    chart_buffer.seek(0)

    # 将图表插入到 Word 文档中
    doc.add_heading('常见协议统计', level=2)
    doc.add_picture(chart_buffer, width=Inches(5.0))

#额外
#-----------------------------流量分析界面生成分析报告----------------

@app.route("/downloadf/", methods=["POST", "GET"])
def downloadf():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
    
        
        time_flow_dict = time_flow(PCAPS)
        host_ip = get_host_ip(PCAPS)
        data_flow_dict = data_flow(PCAPS, host_ip)
        data_ip_dict = data_in_out_ip(PCAPS, host_ip)
        proto_flow_dict = proto_flow(PCAPS)
        most_flow_dict = most_flow_statistic(PCAPS, PD)
        most_flow_dict = sorted(
            most_flow_dict.items(), key=lambda d: d[1], reverse=True
        )
        if len(most_flow_dict) > 10:
            most_flow_dict = most_flow_dict[0:10]
        most_flow_key = [key for key, _ in most_flow_dict]

        # 创建一个新的Word文档
        doc = Document()
        doc.add_heading('数据流分析统计', level=1)


        

        # 将分析结果添加到文档
        #doc.add_heading('时间流量分析', level=2)
        # for key, value in time_flow_dict.items():
        #     doc.add_paragraph(f"{key}: {value}")
        #     # 生成时间流量曲线图
        time_flow_chart = generate_time_flow_chart(time_flow_dict)

        # 将图表保存到 BytesIO 对象中
        chart_buffer = BytesIO()
        canvas = FigureCanvas(time_flow_chart)
        canvas.print_png(chart_buffer)

        # 将图表插入到Word文档中
        doc.add_heading('时间流量曲线图', level=2)
        doc.add_picture(chart_buffer, width=Inches(5.0))


        doc.add_heading('数据流分析', level=2)
        for key, value in data_flow_dict.items():
            doc.add_paragraph(f"{key}: {value}")
                # 生成数据流入流出饼状图
        data_flow_pie_chart = generate_pie_chart(data_flow_dict)

        # 将饼状图保存到 BytesIO 对象中
        pie_chart_buffer = BytesIO()
        data_flow_pie_chart.savefig(pie_chart_buffer, format='png')
        pie_chart_buffer.seek(0)

        # 将饼状图插入到 Word 文档中
        doc.add_heading('数据流入流出统计饼状图', level=2)
        doc.add_picture(pie_chart_buffer, width=Inches(5.0))
        

        doc.add_heading('IP流量分析', level=2)
        for key, value in data_ip_dict.items():
            doc.add_paragraph(f"{key}: {value}")
        add_bar_chart_to_word(doc, data_ip_dict, 'Number of incoming IP traffic packets graph')
        add_bar_chart_to_word_1(doc, data_ip_dict, 'Figure of the number of outgoing IP traffic packets')
        add_bar_chart_to_word_2(doc, data_ip_dict, 'Total incoming IP traffic graph')
        add_bar_chart_to_word_3(doc, data_ip_dict, 'Total outgoing IP traffic graph')
        

        doc.add_heading('协议流量分析', level=2)
        for key, value in proto_flow_dict.items():
            doc.add_paragraph(f"{key}: {value}")
        data_dict = common_proto_statistic(PCAPS)
        add_proto_bar_chart_to_word(doc, data_dict.values(), '常见协议统计', '#87cefa')
        
        # doc.add_heading('Top 10 流量分析', level=2)
        # for key, value in most_flow_dict:
        #     doc.add_paragraph(f"{key}: {value}")

        # 将文档保存到 BytesIO 对象中
        doc_buffer = BytesIO()
        doc.save(doc_buffer)

        # 将 BytesIO 对象的内容发送给用户
        doc_buffer.seek(0)
        return send_file(doc_buffer, as_attachment=True, download_name='flow_analysis_report.docx')


# 访问地图


@app.route("/ipmap/", methods=["POST", "GET"])
def ipmap():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        myip = getmyip()
        if myip:
            host_ip = get_host_ip(PCAPS)
            ipdata = get_ipmap(PCAPS, host_ip)
            geo_dict = ipdata[0]
            ip_value_list = ipdata[1]
            myip_geo = get_geo(myip)
            ip_value_list = [
                (list(d.keys())[0], list(d.values())[0]) for d in ip_value_list
            ]
            print(ip_value_list)
            print(geo_dict)
            return render_template(
                "./dataanalyzer/ipmap.html",
                geo_data=geo_dict,
                ip_value=ip_value_list,
                mygeo=myip_geo,
            )
        else:
            return render_template("./error/neterror.html")

# ----------------------------------------------数据提取页面---------------------------------------------

# Web数据


@app.route("/webdata/", methods=["POST", "GET"])
def webdata():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        host_ip = get_host_ip(PCAPS)
        webdata_list = web_data(PCAPS, host_ip)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if dataid:
            return webdata_list[int(dataid) - 1]["data"].replace("\r\n", "<br>")
        else:

            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(webdata_list), per_page=per_page)
            # pcaps_list = list(webdata_list.values())  # 将字典的值转换为列表
            pcaps_paginated = webdata_list[(page-1)*per_page:page*per_page]

            return render_template("./dataextract/webdata.html", webdata=pcaps_paginated, pagination=pagination)

# Mail数据


@app.route("/maildata/", methods=["POST", "GET"])
def maildata():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        filename = request.args.get("filename")
        datatype = request.args.get("datatype")
        host_ip = get_host_ip(PCAPS)
        mailata_list = mail_data(PCAPS, host_ip)
        filepath = app.config["FILE_FOLDER"] + "Mail/"
        filepath = os.path.abspath(filepath) + "/"


        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15


        if datatype == "raw":
            raw_data = mailata_list[int(dataid) - 1]["data"]
            with open(filepath + "raw_data.txt", "w", encoding="UTF-8") as f:
                f.write(raw_data)
            return send_from_directory(filepath, "raw_data.txt", as_attachment=True)
        if filename and dataid:
            filename_ = (hashlib.md5(filename.encode("UTF-8")
                                     ).hexdigest() + "." + filename.split(".")[-1])
            attachs_dict = mailata_list[int(
                dataid) - 1]["parse_data"]["attachs_dict"]
            mode = "wb"
            encoding = None
            if isinstance(attachs_dict[filename], str):
                mode = "w"
                encoding = "UTF-8"
            elif isinstance(attachs_dict[filename], bytes):
                mode = "wb"
                encoding = None
            with open(os.path.join(filepath, filename), mode, encoding=encoding) as f:
                f.write(attachs_dict[filename])
            return send_from_directory(filepath, filename_, as_attachment=True)
        if dataid:
            # return mailata_list[int(dataid)-1]['data'].replace('\r\n',
            # '<br>')
            maildata = mailata_list[int(dataid) - 1]["parse_data"]
            return render_template(
                "./dataextract/mailparsedata.html", maildata=maildata, dataid=dataid
            )
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(mailata_list), per_page=per_page)
            pcaps_paginated = mailata_list[(page-1)*per_page:page*per_page]

            return render_template("./dataextract/maildata.html", maildata=pcaps_paginated, pagination=pagination)

# FTP数据


@app.route("/ftpdata/", methods=["POST", "GET"])
def ftpdata():
    if PCAPS is None:
        flash("请先上传要分析得数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        host_ip = get_host_ip(PCAPS)
        ftpdata_list = telnet_ftp_data(PCAPS, host_ip, 21)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if dataid:
            return ftpdata_list[int(dataid) - 1]["data"].replace("\r\n", "<br>")
        else:
            
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(ftpdata_list), per_page=per_page)
            pcaps_paginated = ftpdata_list[(page-1)*per_page:page*per_page]
            return render_template("./dataextract/ftpdata.html", ftpdata=pcaps_paginated, pagination=pagination)

# Telnet数据


@app.route("/telnetdata/", methods=["POST", "GET"])
def telnetdata():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        host_ip = get_host_ip(PCAPS)
        telnetdata_list = telnet_ftp_data(PCAPS, host_ip, 23)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if dataid:
            return telnetdata_list[int(dataid) - 1]["data"].replace("\r\n", "<br>")
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(telnetdata_list), per_page=per_page)
            pcaps_paginated = telnetdata_list[(page-1)*per_page:page*per_page]

            return render_template(
                "./dataextract/telnetdata.html", telnetdata=pcaps_paginated, pagination=pagination
            )


# 客户端信息


@app.route("/clientinfo/", methods=["POST", "GET"])
def clientinfo():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        clientinfo_list = client_info(PCAPS)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # 转换为支持分页的数据结构
        pagination = Pagination(page=page, total=len(clientinfo_list), per_page=per_page)
        pcaps_paginated = clientinfo_list[(page-1)*per_page:page*per_page]

        return render_template(
            "./dataextract/clientinfo.html", clientinfos=pcaps_paginated, pagination=pagination
        )


# 敏感数据


@app.route("/sendata/", methods=["POST", "GET"])
def sendata():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        host_ip = get_host_ip(PCAPS)
        sendata_list = sen_data(PCAPS, host_ip)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if dataid:
            return sendata_list[int(dataid) - 1]["data"].replace("\r\n", "<br>")
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(sendata_list), per_page=per_page)
            pcaps_paginated = sendata_list[(page-1)*per_page:page*per_page]
            return render_template("./dataextract/sendata.html", sendata=pcaps_paginated, pagination=pagination)


# ----------------------------------------------一异常信息页面---------------------------------------------

# 异常数据


@app.route("/exceptinfo/", methods=["POST", "GET"])
def exceptinfo():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        dataid = request.args.get("id")
        host_ip = get_host_ip(PCAPS)
        warning_list = exception_warning(PCAPS, host_ip)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if dataid:
            if warning_list[int(dataid) - 1]["data"]:
                return warning_list[int(dataid) - 1]["data"].replace("\r\n", "<br>")
            else:
                return "<center><h3>无相关数据包详情</h3></center>"
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(warning_list), per_page=per_page)
            pcaps_paginated = warning_list[(page-1)*per_page:page*per_page]
            return render_template("./exceptions/exception.html", warning=pcaps_paginated, pagination=pagination)


# ----------------------------------------------文件提取---------------------------------------------
# WEB文件提取


@app.route("/webfile/", methods=["POST", "GET"])
def webfile():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        host_ip = get_host_ip(PCAPS)
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        filepath = app.config["FILE_FOLDER"] + "Web/"
        filepath = os.path.abspath(filepath) + "/"
        if not os.path.exists(filepath):
            os.makedirs(filepath)
        web_list = web_file(PCAPS, host_ip, filepath)
        file_dict = dict()
        for web in web_list:
            file_dict[os.path.split(web["filename"])[-1]] = web["filename"]
        file = request.args.get("file")
        if file in file_dict:
            filename = (hashlib.md5(file.encode("UTF-8")
                                    ).hexdigest() + "." + file.split(".")[-1])

            os.rename(os.path.join(filepath, file),
                      os.path.join(filepath, filename))
            return send_from_directory(filepath, filename, as_attachment=True)
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(web_list), per_page=per_page)
            pcaps_paginated = web_list[(page-1)*per_page:page*per_page]
            return render_template("./fileextract/webfile.html", web_list=pcaps_paginated, pagination=pagination)


# Mail文件提取


@app.route("/mailfile/", methods=["POST", "GET"])
def mailfile():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        host_ip = get_host_ip(PCAPS)

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        filepath = app.config["FILE_FOLDER"] + "Mail/"
        filepath = os.path.abspath(filepath) + "/"
        if not os.path.exists(filepath):
            os.makedirs(filepath)
        mail_list = mail_file(PCAPS, host_ip, filepath)
        file_dict = dict()
        for mail in mail_list:
            file_dict[os.path.split(mail["filename"])[-1]] = mail["filename"]
        file = request.args.get("file")
        if file in file_dict:
            filename = (hashlib.md5(file.encode("UTF-8")
                                    ).hexdigest() + "." + file.split(".")[-1])

            os.rename(os.path.join(filepath, file),
                      os.path.join(filepath, filename))
            return send_from_directory(filepath, filename, as_attachment=True)
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(mail_list), per_page=per_page)
            pcaps_paginated = mail_list[(page-1)*per_page:page*per_page]
            return render_template("./fileextract/mailfile.html", mail_list=pcaps_paginated, pagination=pagination)


# FTP文件提取
@app.route("/ftpfile/", methods=["POST", "GET"])
def ftpfile():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        host_ip = get_host_ip(PCAPS)
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15
        filepath = app.config["FILE_FOLDER"] + "FTP/"
        filepath = os.path.abspath(filepath) + "/"
        if not os.path.exists(filepath):
            os.makedirs(filepath)
        ftp_list = ftp_file(PCAPS, host_ip, filepath)
        file_dict = dict()
        for ftp in ftp_list:
            file_dict[os.path.split(ftp["filename"])[-1]] = ftp["filename"]
        file = request.args.get("file")
        if file in file_dict:
            filename = (hashlib.md5(file.encode("UTF-8")
                                    ).hexdigest() + "." + file.split(".")[-1])
            os.rename(os.path.join(filepath, file),
                      os.path.join(filepath, filename))
            return send_from_directory(filepath, filename, as_attachment=True)
        else:
            # 转换为支持分页的数据结构
            pagination = Pagination(page=page, total=len(ftp_list), per_page=per_page)
            pcaps_paginated = ftp_list[(page-1)*per_page:page*per_page]
            return render_template("./fileextract/ftpfile.html", ftp_list=pcaps_paginated, pagination=pagination)


# 所有二进制文件提取


@app.route("/allfile/", methods=["POST", "GET"])
def allfile():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        filepath = app.config["FILE_FOLDER"] + "All/"
        filepath = os.path.abspath(filepath) + "/"
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        if not os.path.exists(filepath):
            os.makedirs(filepath)
        allfiles_dict = all_files(PCAPS, filepath)
        file = request.args.get("file")
        if file in allfiles_dict:
            filename = (
                hashlib.md5(file.encode("UTF-8")).hexdigest() + "." + file.split(".")[-1])
            os.rename(os.path.join(filepath, file),
                      os.path.join(filepath, filename))
            return send_from_directory(filepath, filename, as_attachment=True)
        else:
            # 转换为支持分页的数据结构
            df = pd.DataFrame(list(allfiles_dict.items()), columns=['Key', 'Value'])  
            print(df)
            # 创建 Paginator 对象
            paginator = Pagination(page=page, per_page=per_page, total=len(df), css_framework='bootstrap4')

            # 将 DataFrame 转换为支持分页的数据结构
            pcaps_list = df.to_dict(orient='records')
            pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

            # pagination = Pagination(page=page, total=len(allfiles_dict), per_page=per_page)
            # pcaps_list = list(allfiles_dict)  # 将字典的值转换为列表
            # print(1,allfiles_dict,pcaps_list)
            # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]
            # print(2,pcaps_paginated)
            return render_template(
                "./fileextract/allfile.html", allfiles_dict=pcaps_paginated, pagination=paginator
            )


# ----------------------------------------------错误处理页面---------------------------------------------


@app.errorhandler(404)
def internal_error404(error):
    return render_template("./error/404.html"), 404


@ app.errorhandler(500)
def internal_error500(error):
    return render_template("./error/500.html"), 500


# ---------------------------------------------NEW PAGE----------------------

# 加密流量识别

@app.route("/behavioranalysis1/", methods=["POST", "GET"])
def isVPN():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        global flow_df
        flow_df = predict_csv(PCAP_PATH, 5)
        flow_df = [] if flow_df is None else flow_df
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # # 转换为支持分页的数据结构
        # pagination = Pagination(page=page, total=len(flow_df), per_page=per_page)
        # pcaps_list = list(flow_df.values())  # 将字典的值转换为列表
        # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        # 创建 Paginator 对象
        paginator = Pagination(page=page, per_page=per_page, total=len(flow_df), css_framework='bootstrap4')

        # 将 DataFrame 转换为支持分页的数据结构
        pcaps_list = flow_df.to_dict(orient='records')
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        print(pcaps_paginated)

        return render_template(
            "./behavior/behavioranalysis1.html", flow_df=pcaps_paginated, pagination=paginator
        )

# 加密流量分类

@app.route("/behavioranalysis2/", methods=["POST", "GET"])
def VPNfield():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        flow_df = predict_csv(PCAP_PATH, 6)
        flow_df = [] if flow_df is None else flow_df

        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # # 转换为支持分页的数据结构
        # pagination = Pagination(page=page, total=len(flow_df), per_page=per_page)
        # pcaps_list = list(flow_df.values())  # 将字典的值转换为列表
        # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        # 创建 Paginator 对象
        paginator = Pagination(page=page, per_page=per_page, total=len(flow_df), css_framework='bootstrap4')

        # 将 DataFrame 转换为支持分页的数据结构
        pcaps_list = flow_df.to_dict(orient='records')
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]


        return render_template(
            "./behavior/behavioranalysis2.html", flow_df=pcaps_paginated, pagination=paginator
        )

# ---二分类任务---

# 恶意加密流量识别

@app.route("/behavioranalysis/", methods=["POST", "GET"])
def behavioranalysis():
    # 最后补回来
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        flow_df = predict_csv(PCAP_PATH, 4)#是否为VPN
        # 分析结果 csvfile
        # TODO 显示
        flow_df = [] if flow_df is None else flow_df


        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # # 转换为支持分页的数据结构
        # pagination = Pagination(page=page, total=len(flow_df), per_page=per_page)
        # pcaps_list = list(flow_df.values())  # 将字典的值转换为列表
        # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        # 创建 Paginator 对象
        paginator = Pagination(page=page, per_page=per_page, total=len(flow_df), css_framework='bootstrap4')

        # 将 DataFrame 转换为支持分页的数据结构
        pcaps_list = flow_df.to_dict(orient='records')
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]
        
        return render_template(
            "./behavior/behavioranalysis.html", flow_df=pcaps_paginated, pagination=paginator
        )
# ---恶意流量四分类任务---

# 恶意加密流量分类

@app.route("/malwareTypeClass/", methods=["POST", "GET"])
def malwareTypeClass():
    # 最后补回来
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        global flow_df3
        if flow_df3 is None:
            flow_df3 = predict_csv(PCAP_PATH, 3)#是否为4种恶意加密流量
        # 分析结果 csvfile
        # TODO 显示
        flow_df3 = [] if flow_df3 is None else flow_df3
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # # 转换为支持分页的数据结构
        # pagination = Pagination(page=page, total=len(flow_df3), per_page=per_page)
        # pcaps_list = list(flow_df.values())  # 将字典的值转换为列表
        # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        # 创建 Paginator 对象
        paginator = Pagination(page=page, per_page=per_page, total=len(flow_df3), css_framework='bootstrap4')

        # 将 DataFrame 转换为支持分页的数据结构
        pcaps_list = flow_df3.to_dict(orient='records')
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        return render_template(
            "./behavior/malwareTypeClass.html", flow_df=pcaps_paginated, pagination=paginator
        )

# 恶意加密流量类别

@app.route("/malwareMultiType/", methods=["POST", "GET"])
def malwareMultiType():
    if PCAPS is None:
        flash("请先上传要分析的数据包!")
        return redirect(url_for("upload"))
    else:
        flow_df = predict_csv(PCAP_PATH, 5)
        flow_df = [] if flow_df is None else flow_df
        page = request.args.get("page", type=int, default=None)

        if page == None:
            page = 1
        per_page = 15

        # # 转换为支持分页的数据结构
        # pagination = Pagination(page=page, total=len(flow_df3), per_page=per_page)
        # pcaps_list = list(flow_df.values())  # 将字典的值转换为列表
        # pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        # 创建 Paginator 对象
        paginator = Pagination(page=page, per_page=per_page, total=len(flow_df), css_framework='bootstrap4')

        # 将 DataFrame 转换为支持分页的数据结构
        pcaps_list = flow_df.to_dict(orient='records')
        pcaps_paginated = pcaps_list[(page-1)*per_page:page*per_page]

        return render_template(
            "./behavior/malwareMultiType.html", flow_df=pcaps_paginated, pagination=paginator
        )


@app.route("/configsys", methods=["POST", "GET"])

def configsys():
    if request.method == "GET":
        return render_template("./behavior/configSys.html")

    elif request.method == "POST":
        data = request.get_json()
        NUM = data.get('NUM')
        M1_LIMIT = data.get('M1_LIMIT')

        MODE_3_LOSS = data.get('MODE_3_LOSS')
        MODE_3_LEARNINGRATE = float(data.get('MODE_3_LEARNINGRATE'))
        MODE_3_METRICS = data.get('MODE_3_METRICS')
        MODE_3_BATCH = int(data.get('MODE_3_BATCH'))

        MODE_4_LOSS = data.get('MODE_4_LOSS')
        MODE_4_LEARNINGRATE = float(data.get('MODE_4_LEARNINGRATE'))
        MODE_4_METRICS = data.get('MODE_4_METRICS')
        MODE_4_BATCH = int(data.get('MODE_4_BATCH'))

        MODE_5_LOSS = data.get('MODE_5_LOSS')
        MODE_5_LEARNINGRATE = float(data.get('MODE_5_LEARNINGRATE'))
        MODE_5_METRICS = data.get('MODE_5_METRICS')
        MODE_5_BATCH = int(data.get('MODE_5_BATCH'))
        with open('./config.py', 'r') as f:
            lines = f.readlines()
        for i, line in enumerate(lines):
            if line.startswith('NUM'):
                a = NUM
                lines[i] = str("NUM ="+ str(a)+'\n')
        with open('./config.py', 'w') as f:
            f.writelines(lines)       
#------------------------------------------------------------------
        with open('./config.py', 'r') as f:
            lines = f.readlines()
        for i, line in enumerate(lines):
            if line.startswith('M1_LIMIT'):
                a = M1_LIMIT
                lines[i] = str("M1_LIMIT ="+ str(a)+'\n')
        with open('./config.py', 'w') as f:
            f.writelines(lines)
#------------------------------------------------------------------
        with open('./config.py', 'r') as f:
            lines = f.readlines()
        for i, line in enumerate(lines):
            if line.startswith('MODE_3'):
                rate = MODE_3_LEARNINGRATE
                batch = MODE_3_BATCH
                loss = MODE_3_LOSS
                acc = MODE_3_METRICS
                mode = [loss,rate,acc,batch]
                lines[i] = str("MODE_3 ="+ str(mode)+'\n')
        with open('./config.py', 'w') as f:
            f.writelines(lines)     
#------------------------------------------------------------------
        with open('./config.py', 'r') as f:
            lines = f.readlines()
        for i, line in enumerate(lines):
            if line.startswith('MODE_4'):
                rate = MODE_4_LEARNINGRATE
                batch = MODE_4_BATCH
                loss = MODE_4_LOSS
                acc = MODE_4_METRICS
                mode = [loss,rate,acc,batch]
                lines[i] = str("MODE_4 ="+ str(mode)+'\n')
        with open('./config.py', 'w') as f:
            f.writelines(lines)                            
#------------------------------------------------------------------
        with open('./config.py', 'r') as f:
            lines = f.readlines()
        for i, line in enumerate(lines):
            if line.startswith('MODE_5'):
                rate = MODE_5_LEARNINGRATE
                batch = MODE_5_BATCH
                loss = MODE_5_LOSS
                acc = MODE_5_METRICS
                mode = [loss,rate,acc,batch]
                lines[i] = str("MODE_5 ="+ str(mode)+'\n')
        with open('./config.py', 'w') as f:
            f.writelines(lines) 
#------------------------------------------------------------------
            
        return jsonify(code="200")

#额外
@app.route('/export')
def export_data():
    # 获取导出格式，这里使用 request.args.get('exportFormat', 'xlsx') 获取前端传递的参数
    export_format = request.args.get('format', 'xml')
    global flow_df  
    # 生成 flow_df 数据，flow_df 类型为 <class 'pandas.core.frame.DataFrame'>
    # flow_df = predict_csv(PCAP_PATH, 5)
   # 将列名中以数字开头的部分替换为有效的标签名称
    flow_df.columns = ["源IP地址", "源端口", "目的IP地址", "目的端口","predicted_labels"]
    #print(flow_df)
    # print(export_format)
   # 导出数据到指定格式
    # 导出数据到指定格式
    if export_format == 'xml':
        response = make_response(flow_df.to_xml())
        response.headers['Content-Type'] = 'application/xml'
        response.headers['Content-Disposition'] = 'attachment; filename=exported_data.xml'
        response.data = response.data.replace(b"<?xml version='1.0' encoding='utf-8'?>", b"<?xml-stylesheet type='text/css' href='style.css'?>")
    elif export_format == 'xls':
            # 使用 xlwt 库将 DataFrame 写入 BytesIO
        excel_data = io.BytesIO()
        
        # 创建一个Workbook对象
        workbook = xlwt.Workbook()
        
        # 添加一个Sheet，并获取它
        sheet = workbook.add_sheet('Sheet1')
        
        # 将DataFrame写入Sheet
        for i, col in enumerate(flow_df.columns):
            sheet.write(0, i, col)  # 写入列名
            for j, value in enumerate(flow_df[col]):
                sheet.write(j + 1, i, value)  # 写入数据
        
        # 保存Workbook到BytesIO
        workbook.save(excel_data)
        
        response = make_response(excel_data.getvalue())
        response.headers['Content-Type'] = 'application/vnd.ms-excel'
        response.headers['Content-Disposition'] = 'attachment; filename=exported_data.xls'
    elif export_format == 'xlsx':
        # 使用 openpyxl 库将 DataFrame 写入 BytesIO
        excel_data = io.BytesIO()
        flow_df.to_excel(excel_data, index=False, sheet_name='Sheet1', engine='openpyxl')
        
        response = make_response(excel_data.getvalue())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        response.headers['Content-Disposition'] = 'attachment; filename=exported_data.xlsx'
    else:
        # 处理未知的导出格式，你可以根据需求进行处理
        return "不支持的导出格式", 400

    return response

#额外
def generate_unique_filename(filename):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    random_string = os.urandom(8).hex()
    _, extension = os.path.splitext(filename)
    new_filename = f"{timestamp}_{random_string}{extension}"
    return new_filename

#额外
#图像隐写分析
def is_grayscale(image_path):
    # 使用Pillow库打开图像
    with Image.open(image_path) as img:
        # 检查通道数
        return img.mode == 'L'
@app.route("/imageSteganalysis/", methods=["POST", "GET"])
def imageSteganalysis():
    result = None
    error_message=None
    if 'image' in request.files:
        image_file = request.files['image']
        print(type(image_file))
        
        # 生成新的文件名
        new_filename = generate_unique_filename(image_file.filename)
        
        # 保存文件到服务器
        image_path = os.path.join('uploads/', new_filename)
        # print(image_path)
        image_file.save(image_path)
        if not is_grayscale(image_path):
            error_message = '错误：只允许使用灰度图像。'
        else:
            model_path = "./XuNet/xu/xu_0.4_suniward_epoch_159_best_acc_val64.95_test80.17.pkl"
            detector = StegoDetector(model_path)
            # image_path = "/root/pcap_analysis/XuNet/5001.pgm"
            stego_probability = detector.detect_stego_probability(image_path)
            result = f'The probability of the stego sample: {stego_probability}'
            print(f'The probability of the stego sample: {stego_probability}')
            # print(f"成功保存，文件名为：{new_filename}")
        
    return render_template(
        "./imageSteganalysis/imageSteganalysis.html",result=result,error_message=error_message
        )